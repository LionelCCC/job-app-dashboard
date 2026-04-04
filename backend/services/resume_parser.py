"""
Resume parsing service for the Job Application Pipeline Dashboard.
Supports PDF (via PyMuPDF/fitz), DOCX (via python-docx), and LaTeX (.tex).
"""

import re
import logging
import os
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Optional import wrappers – handle gracefully if libs are not installed
# ---------------------------------------------------------------------------

try:
    import fitz  # PyMuPDF
    _FITZ_AVAILABLE = True
except ImportError:
    _FITZ_AVAILABLE = False
    logger.warning("PyMuPDF (fitz) not installed – PDF parsing disabled")

try:
    from docx import Document as DocxDocument
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False
    logger.warning("python-docx not installed – DOCX parsing disabled")


# ---------------------------------------------------------------------------
# Low-level text extraction
# ---------------------------------------------------------------------------

def _extract_text_pdf(file_path: str) -> str:
    """Extract full text from a PDF using PyMuPDF."""
    if not _FITZ_AVAILABLE:
        raise RuntimeError("PyMuPDF is not installed. Run: pip install PyMuPDF")
    text_parts: list[str] = []
    doc = fitz.open(file_path)
    try:
        for page in doc:
            text_parts.append(page.get_text("text"))
    finally:
        doc.close()
    return "\n".join(text_parts)


def _extract_text_docx(file_path: str) -> str:
    """Extract full text from a DOCX file using python-docx."""
    if not _DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")
    doc = DocxDocument(file_path)
    paragraphs = [para.text for para in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.append(cell.text)
    return "\n".join(paragraphs)


def _extract_text_latex(file_path: str) -> str:
    """
    Extract plain text from a LaTeX (.tex) file by stripping LaTeX commands.

    Strategy:
    1. Read raw .tex source
    2. Remove comments (% to end of line)
    3. Remove common structural commands: \\begin{}, \\end{}, \\documentclass, etc.
    4. Replace \\section, \\subsection, etc. with their text arguments
    5. Remove remaining \\command[opt]{arg} patterns, keeping the arg text where sensible
    6. Clean up whitespace
    """
    with open(file_path, encoding="utf-8", errors="replace") as f:
        src = f.read()

    # Remove LaTeX comments
    src = re.sub(r"%.*", "", src)

    # Remove document preamble (everything before \begin{document})
    begin_doc = src.find(r"\begin{document}")
    if begin_doc != -1:
        src = src[begin_doc + len(r"\begin{document}"):]

    # Remove \end{document} and beyond
    end_doc = src.find(r"\end{document}")
    if end_doc != -1:
        src = src[:end_doc]

    # Extract text from section-type commands (keep the argument)
    src = re.sub(r"\\(?:section|subsection|subsubsection|paragraph|textbf|textit|emph|underline|textsc|textrm)\*?\{([^}]*)\}", r"\1", src)

    # Remove \begin{} and \end{} environment markers
    src = re.sub(r"\\(?:begin|end)\{[^}]*\}", "", src)

    # Remove common formatting commands that take an argument (keep content)
    src = re.sub(r"\\(?:item|label|ref|cite|href|url|footnote|caption)\{([^}]*)\}", r"\1", src)
    src = re.sub(r"\\(?:href|hyperref)\[[^\]]*\]\{([^}]*)\}", r"\1", src)

    # Remove commands with optional args but keep mandatory arg
    src = re.sub(r"\\[a-zA-Z]+\[[^\]]*\]\{([^}]*)\}", r"\1", src)

    # Remove remaining \command{arg} patterns — drop command, keep arg
    src = re.sub(r"\\[a-zA-Z]+\{([^}]*)\}", r"\1", src)

    # Remove lone \command (no args)
    src = re.sub(r"\\[a-zA-Z]+\*?", "", src)

    # Remove remaining braces, backslashes, and LaTeX special chars
    src = re.sub(r"[{}\\]", " ", src)

    # Normalize whitespace
    src = re.sub(r"[ \t]+", " ", src)
    src = re.sub(r"\n{3,}", "\n\n", src)

    return src.strip()


def _extract_text(file_path: str) -> str:
    """Dispatch to the appropriate extractor based on file extension."""
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return _extract_text_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return _extract_text_docx(file_path)
    elif ext in (".tex", ".latex"):
        return _extract_text_latex(file_path)
    else:
        raise ValueError(
            f"Unsupported file type: {ext}. Supported: PDF, DOCX, LaTeX (.tex)."
        )


# ---------------------------------------------------------------------------
# Regex-based field extractors
# ---------------------------------------------------------------------------

_EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}")
_PHONE_RE = re.compile(
    r"(\+?1[\s\-.]?)?\(?\d{3}\)?[\s\-.]?\d{3}[\s\-.]?\d{4}"
)
_LINKEDIN_RE = re.compile(
    r"(?:linkedin\.com/in/|linkedin\.com/pub/)([A-Za-z0-9\-_%]+)", re.IGNORECASE
)
_GITHUB_RE = re.compile(r"(?:github\.com/)([A-Za-z0-9\-_%]+)", re.IGNORECASE)


def _extract_contact_info(text: str) -> dict:
    """Return a dict with name, email, phone, linkedin, github."""
    contact: dict = {"name": "", "email": "", "phone": "", "linkedin": "", "github": ""}

    emails = _EMAIL_RE.findall(text)
    if emails:
        contact["email"] = emails[0]

    phones = _PHONE_RE.findall(text)
    if phones:
        raw = phones[0] if isinstance(phones[0], str) else "".join(phones[0])
        contact["phone"] = re.sub(r"\s+", " ", raw).strip()

    linkedin_match = _LINKEDIN_RE.search(text)
    if linkedin_match:
        contact["linkedin"] = f"linkedin.com/in/{linkedin_match.group(1)}"

    github_match = _GITHUB_RE.search(text)
    if github_match:
        contact["github"] = f"github.com/{github_match.group(1)}"

    # Name heuristic: first non-empty line that looks like a proper name
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        if (
            "@" in line
            or re.search(r"\d{3}", line)
            or len(line) > 60
            or line.isupper()
            or line.lower().startswith(
                ("summary", "objective", "experience", "education", "skills")
            )
        ):
            continue
        words = line.split()
        if 2 <= len(words) <= 4 and all(w[0].isupper() for w in words if w):
            contact["name"] = line
            break

    return contact


_SECTION_HEADERS = {
    "experience": re.compile(
        r"^\s*(work\s*)?experience|employment\s*history|professional\s*background",
        re.IGNORECASE,
    ),
    "education": re.compile(r"^\s*education|academic\s*background", re.IGNORECASE),
    "skills": re.compile(
        r"^\s*skills|technical\s*skills|core\s*competencies|technologies",
        re.IGNORECASE,
    ),
    "certifications": re.compile(
        r"^\s*certifications?|licenses?|credentials?", re.IGNORECASE
    ),
    "summary": re.compile(r"^\s*summary|objective|profile|about\s*me", re.IGNORECASE),
    "projects": re.compile(r"^\s*projects?", re.IGNORECASE),
}


def _split_sections(text: str) -> dict[str, str]:
    lines = text.splitlines()
    sections: dict[str, list[str]] = {"header": []}
    current_section = "header"

    for line in lines:
        matched = False
        for section_name, pattern in _SECTION_HEADERS.items():
            if pattern.match(line.strip()):
                current_section = section_name
                sections.setdefault(current_section, [])
                matched = True
                break
        if not matched:
            sections.setdefault(current_section, [])
            sections[current_section].append(line)

    return {k: "\n".join(v).strip() for k, v in sections.items()}


def _extract_skills(skills_text: str, full_text: str) -> list[str]:
    skills: set[str] = set()

    if skills_text:
        for token in re.split(r"[,|•·\n\t/]+", skills_text):
            token = token.strip(" -•·\t\r")
            if 1 < len(token) < 50:
                skills.add(token)

    known_skills = [
        "Python", "Java", "Scala", "Kotlin", "JavaScript", "TypeScript", "Go", "Golang",
        "C++", "C#", "Ruby", "Rust", "Swift", "R", "MATLAB", "SQL", "Bash", "Shell",
        "FastAPI", "Django", "Flask", "Spring Boot", "React", "Vue", "Angular", "Node.js",
        "TensorFlow", "PyTorch", "Keras", "Scikit-learn", "XGBoost", "LightGBM",
        "Hugging Face", "LangChain", "OpenAI", "Spark", "PySpark", "Flink", "Kafka",
        "Airflow", "dbt", "Dagster", "Prefect", "MLflow", "Kubeflow", "Ray",
        "AWS", "GCP", "Azure", "Docker", "Kubernetes", "Terraform", "Ansible",
        "GitHub Actions", "CI/CD", "Jenkins", "GitLab CI",
        "PostgreSQL", "MySQL", "SQLite", "MongoDB", "Redis", "Elasticsearch",
        "Cassandra", "DynamoDB", "BigQuery", "Snowflake", "Redshift", "Databricks",
        "Tableau", "Power BI", "Looker", "Grafana", "Metabase",
        "REST API", "GraphQL", "gRPC", "Microservices", "Agile", "Scrum",
        "Git", "Linux", "Pandas", "NumPy", "Matplotlib", "Seaborn",
    ]
    full_lower = full_text.lower()
    for skill in known_skills:
        if skill.lower() in full_lower:
            skills.add(skill)

    return sorted(skills)


def _extract_experience(experience_text: str) -> list[dict]:
    if not experience_text:
        return []

    entries: list[dict] = []
    date_pattern = re.compile(
        r"(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* \d{4}"
        r"|\d{4}\s*[-–]\s*(\d{4}|present|current|now)",
        re.IGNORECASE,
    )

    current: Optional[dict] = None
    body_lines: list[str] = []

    def _flush_current():
        if current is not None:
            current["description"] = " ".join(body_lines).strip()
            entries.append(current)

    for line in experience_text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue

        has_date = bool(date_pattern.search(stripped))
        is_short = len(stripped) < 120

        if has_date and is_short:
            _flush_current()
            body_lines = []
            parts = re.split(r"\s+at\s+|\s*[|·—]\s*|\s*,\s*", stripped, maxsplit=2)
            role = parts[0].strip() if parts else stripped
            company = parts[1].strip() if len(parts) > 1 else ""
            duration_match = date_pattern.search(stripped)
            duration = duration_match.group(0) if duration_match else ""
            current = {"role": role, "company": company, "duration": duration, "description": ""}
        else:
            body_lines.append(stripped)

    _flush_current()

    if not entries and experience_text.strip():
        entries.append({
            "role": "", "company": "", "duration": "",
            "description": experience_text.strip()[:1000],
        })

    return entries


def _extract_education(education_text: str) -> list[dict]:
    if not education_text:
        return []

    entries: list[dict] = []
    year_re = re.compile(r"\b(19|20)\d{2}\b")
    degree_keywords = re.compile(
        r"\b(bachelor|b\.?s\.?|b\.?a\.?|master|m\.?s\.?|m\.?a\.?|"
        r"ph\.?d\.?|doctor|associate|mba|b\.?eng\.?|m\.?eng\.?)\b",
        re.IGNORECASE,
    )

    lines = [li.strip() for li in education_text.splitlines() if li.strip()]
    i = 0
    while i < len(lines):
        line = lines[i]
        year_match = year_re.search(line)
        degree_match = degree_keywords.search(line)
        if year_match or degree_match:
            year = year_match.group(0) if year_match else ""
            degree = degree_match.group(0) if degree_match else ""
            institution = lines[i + 1] if i + 1 < len(lines) else ""
            entries.append({"institution": institution, "degree": degree, "year": year})
            i += 2
        else:
            i += 1

    if not entries and education_text.strip():
        entries.append({
            "institution": education_text.strip()[:200], "degree": "", "year": ""
        })

    return entries


def _extract_certifications(cert_text: str) -> list[str]:
    if not cert_text:
        return []
    certs: list[str] = []
    for line in cert_text.splitlines():
        line = line.strip(" -•·\t\r")
        if len(line) > 3:
            certs.append(line)
    return certs


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def parse_resume(file_path: str) -> dict:
    """
    Parse a PDF, DOCX, or LaTeX resume file into a structured dict.

    Returns
    -------
    dict with keys:
        full_text, contact_info, skills, experience, education,
        certifications, summary
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Resume file not found: {file_path}")

    full_text = _extract_text(file_path)
    sections = _split_sections(full_text)

    return {
        "full_text": full_text,
        "contact_info": _extract_contact_info(full_text),
        "skills": _extract_skills(sections.get("skills", ""), full_text),
        "experience": _extract_experience(sections.get("experience", "")),
        "education": _extract_education(sections.get("education", "")),
        "certifications": _extract_certifications(sections.get("certifications", "")),
        "summary": sections.get("summary", "")[:500],
    }


def categorize_resume(parsed_data: dict) -> str:
    """
    Determine the most appropriate job-type category for a resume.
    Returns one of: SWE, DE, DA, DS, MLE, AIE
    """
    from services.job_scraper import determine_job_type

    skills_text = " ".join(parsed_data.get("skills", []))
    exp_text = " ".join(
        f"{e.get('role', '')} {e.get('description', '')}"
        for e in parsed_data.get("experience", [])
    )
    summary = parsed_data.get("summary", "")
    return determine_job_type("", f"{skills_text} {exp_text} {summary}")
